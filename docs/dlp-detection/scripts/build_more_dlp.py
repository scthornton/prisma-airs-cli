#!/usr/bin/env python3
"""More DLP test vectors:

CONTROLS (to explain the #4 'toxic' finding):
  dlp_ctrl_clean.png         - base image as PNG, NO steg  (false-positive control)
  dlp_ctrl_stego_benign.png  - LSB steg carrying ONLY benign lorem-ipsum, no PII

NEW MODALITIES:
  dlp_doc_sensitive.docx     - synthetic PII in body + hidden white micro-text + core metadata
  dlp_archive.zip            - zip containing payload.txt (archive-recursion test)
"""
import struct, zipfile
import numpy as np
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor

BASE = "dlp_img_base.jpg"
PAYLOAD_LINES = [
    "SSN: 078-05-1120",
    "Visa 4111 1111 1111 1111",
    "AWS_KEY AKIAIOSFODNN7EXAMPLE",
    "AWS_SECRET wJalrXUtnFEMI/K7MDENG/bPxRfiCYEXAMPLEKEY",
    "Email john.public@example.com  Phone (555) 010-0142",
    "Passport X12345678  DOB 1985-07-14",
]
BLOB = "\n".join(PAYLOAD_LINES)
BENIGN = ("Lorem ipsum dolor sit amet, consectetur adipiscing elit. "
          "Curabitur nec quam vitae arcu pretium tincidunt. "
          "This file contains no sensitive data whatsoever.")


def lsb_embed(cover_arr, data: bytes):
    flat = cover_arr.reshape(-1).copy()
    blob = struct.pack(">I", len(data)) + data
    bits = np.unpackbits(np.frombuffer(blob, dtype=np.uint8))
    flat[:bits.size] = (flat[:bits.size] & 0xFE) | bits
    return flat.reshape(cover_arr.shape)


def main():
    base = np.array(Image.open(BASE).convert("RGB"), dtype=np.uint8)

    # control 1: clean PNG, no steg
    Image.fromarray(base, "RGB").save("dlp_ctrl_clean.png")

    # control 2: steg PNG with benign payload only
    Image.fromarray(lsb_embed(base, BENIGN.encode()), "RGB").save(
        "dlp_ctrl_stego_benign.png")

    # docx: visible PII + hidden white micro-text + metadata
    doc = Document()
    doc.add_heading("Quarterly Account Review", 0)
    doc.add_paragraph("The following records were flagged for audit:")
    for line in PAYLOAD_LINES:
        doc.add_paragraph(line)
    # hidden run: white, 1pt
    p = doc.add_paragraph()
    r = p.add_run("HIDDEN " + BLOB.replace("\n", " | "))
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(1)
    cp = doc.core_properties
    cp.author = "John Q. Public (synthetic)"
    cp.comments = "SSN 078-05-1120; Visa 4111 1111 1111 1111; AKIAIOSFODNN7EXAMPLE"
    cp.keywords = "SSN;PAN;AWS;PII;passport;X12345678"
    doc.save("dlp_doc_sensitive.docx")

    # zip containing the payload as a text file
    with open("payload.txt", "w") as f:
        f.write(BLOB + "\n")
    with zipfile.ZipFile("dlp_archive.zip", "w", zipfile.ZIP_DEFLATED) as z:
        z.write("payload.txt")

    for f in ["dlp_ctrl_clean.png", "dlp_ctrl_stego_benign.png",
              "dlp_doc_sensitive.docx", "dlp_archive.zip"]:
        print("Generated:", f)


if __name__ == "__main__":
    main()
